import openpyxl
import pandas as pd
from docx import Document
from docx.shared import Inches
import os
from datetime import datetime
import nltk
from nltk.tokenize import sent_tokenize
from transformers import pipeline
import spacy
import logging
import numpy as np
from sklearn.cluster import KMeans
from sklearn.preprocessing import StandardScaler
import matplotlib.pyplot as plt
import seaborn as sns
from collections import Counter
import warnings
warnings.filterwarnings('ignore')

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

class ExcelAIAnalyzer:
    def __init__(self):
        """Initialize the Excel AI Analyzer with necessary models and configurations."""
        try:
            # Download required NLTK data
            nltk.download('punkt')
            nltk.download('averaged_perceptron_tagger')
            nltk.download('maxent_ne_chunker')
            nltk.download('words')
            
            # Load spaCy model
            self.nlp = spacy.load('en_core_web_sm')
            
            # Initialize summarization pipeline
            self.summarizer = pipeline("summarization", model="facebook/bart-large-cnn")
            
            # Initialize NER pipeline
            self.ner_pipeline = pipeline("ner", model="dbmdz/bert-large-cased-finetuned-conll03-english")
            
            logger.info("Excel AI Analyzer initialized successfully")
        except Exception as e:
            logger.error(f"Error initializing Excel AI Analyzer: {str(e)}")
            raise

    def read_excel_data(self, file_path, sheet_name=None):
        """Read data from Excel file and return as pandas DataFrame."""
        try:
            if sheet_name:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
            else:
                df = pd.read_excel(file_path)
            
            # Convert date columns to datetime
            date_columns = df.select_dtypes(include=['object']).columns
            for col in date_columns:
                try:
                    df[col] = pd.to_datetime(df[col])
                except:
                    continue
                    
            logger.info(f"Successfully read data from {file_path}")
            return df
        except Exception as e:
            logger.error(f"Error reading Excel file: {str(e)}")
            raise

    def analyze_data(self, df):
        """Analyze the data and generate insights."""
        try:
            insights = {
                'summary': self._generate_summary(df),
                'statistics': self._calculate_statistics(df),
                'trends': self._identify_trends(df),
                'surveillance_analysis': self._analyze_surveillance_data(df),
                'risk_assessment': self._assess_risk_factors(df),
                'geographic_analysis': self._analyze_geographic_distribution(df),
                'temporal_analysis': self._analyze_temporal_patterns(df)
            }
            return insights
        except Exception as e:
            logger.error(f"Error analyzing data: {str(e)}")
            raise

    def _analyze_surveillance_data(self, df):
        """Analyze disease surveillance data."""
        try:
            analysis = {}
            
            # Case distribution by diagnosis
            if 'Diagnosis' in df.columns:
                analysis['diagnosis_distribution'] = df['Diagnosis'].value_counts().to_dict()
            
            # Age group analysis
            if 'Age' in df.columns:
                df['Age_Group'] = pd.cut(df['Age'], 
                                       bins=[0, 5, 12, 18, 60, 100],
                                       labels=['0-5', '6-12', '13-18', '19-60', '60+'])
                analysis['age_distribution'] = df['Age_Group'].value_counts().to_dict()
            
            # Gender distribution
            if 'Gender' in df.columns:
                analysis['gender_distribution'] = df['Gender'].value_counts().to_dict()
            
            # Outcome analysis
            if 'Outcome' in df.columns:
                analysis['outcome_distribution'] = df['Outcome'].value_counts().to_dict()
            
            # Lab result analysis
            if 'Lab_Result' in df.columns:
                analysis['lab_result_distribution'] = df['Lab_Result'].value_counts().to_dict()
            
            return analysis
        except Exception as e:
            logger.error(f"Error in surveillance analysis: {str(e)}")
            return {}

    def _assess_risk_factors(self, df):
        """Assess risk factors for disease transmission."""
        try:
            risk_factors = {}
            
            # Age-based risk
            if 'Age' in df.columns:
                risk_factors['age_risk'] = {
                    'high_risk_age_groups': df[df['Age'] < 5]['Age'].count(),
                    'moderate_risk_age_groups': df[(df['Age'] >= 5) & (df['Age'] < 18)]['Age'].count()
                }
            
            # Geographic clustering
            if 'District' in df.columns:
                district_counts = df['District'].value_counts()
                risk_factors['geographic_clusters'] = {
                    'high_incidence_areas': district_counts[district_counts > district_counts.mean()].to_dict()
                }
            
            # Temporal clustering
            if 'Date' in df.columns:
                df['Week'] = df['Date'].dt.isocalendar().week
                weekly_counts = df.groupby('Week').size()
                risk_factors['temporal_clusters'] = {
                    'high_incidence_weeks': weekly_counts[weekly_counts > weekly_counts.mean()].to_dict()
                }
            
            return risk_factors
        except Exception as e:
            logger.error(f"Error in risk assessment: {str(e)}")
            return {}

    def _analyze_geographic_distribution(self, df):
        """Analyze geographic distribution of cases."""
        try:
            geo_analysis = {}
            
            if 'District' in df.columns:
                # District-wise case distribution
                geo_analysis['district_distribution'] = df['District'].value_counts().to_dict()
                
                # District-wise outcome analysis
                if 'Outcome' in df.columns:
                    geo_analysis['district_outcomes'] = df.groupby(['District', 'Outcome']).size().to_dict()
                
                # District-wise lab results
                if 'Lab_Result' in df.columns:
                    geo_analysis['district_lab_results'] = df.groupby(['District', 'Lab_Result']).size().to_dict()
            
            return geo_analysis
        except Exception as e:
            logger.error(f"Error in geographic analysis: {str(e)}")
            return {}

    def _analyze_temporal_patterns(self, df):
        """Analyze temporal patterns in the data."""
        try:
            temporal_analysis = {}
            
            if 'Date' in df.columns:
                # Daily case counts
                df['Date'] = pd.to_datetime(df['Date'])
                daily_cases = df.groupby(df['Date'].dt.date).size()
                temporal_analysis['daily_cases'] = daily_cases.to_dict()
                
                # Weekly trends
                weekly_cases = df.groupby(df['Date'].dt.isocalendar().week).size()
                temporal_analysis['weekly_trends'] = weekly_cases.to_dict()
                
                # Monthly trends
                monthly_cases = df.groupby(df['Date'].dt.month).size()
                temporal_analysis['monthly_trends'] = monthly_cases.to_dict()
            
            return temporal_analysis
        except Exception as e:
            logger.error(f"Error in temporal analysis: {str(e)}")
            return {}

    def generate_report(self, insights, output_path):
        """Generate a comprehensive Word document report with the analysis results."""
        try:
            doc = Document()
            
            # Add title
            doc.add_heading('Disease Surveillance Analysis Report', 0)
            doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            
            # Add executive summary
            doc.add_heading('Executive Summary', level=1)
            doc.add_paragraph(insights['summary'])
            
            # Add surveillance analysis
            doc.add_heading('Surveillance Analysis', level=1)
            surveillance = insights['surveillance_analysis']
            
            if 'diagnosis_distribution' in surveillance:
                doc.add_heading('Case Distribution by Diagnosis', level=2)
                for diagnosis, count in surveillance['diagnosis_distribution'].items():
                    doc.add_paragraph(f'{diagnosis}: {count} cases')
            
            if 'age_distribution' in surveillance:
                doc.add_heading('Age Distribution', level=2)
                for age_group, count in surveillance['age_distribution'].items():
                    doc.add_paragraph(f'{age_group}: {count} cases')
            
            # Add risk assessment
            doc.add_heading('Risk Assessment', level=1)
            risk = insights['risk_assessment']
            
            if 'age_risk' in risk:
                doc.add_heading('Age-based Risk Analysis', level=2)
                doc.add_paragraph(f'High-risk age groups (0-5 years): {risk["age_risk"]["high_risk_age_groups"]} cases')
                doc.add_paragraph(f'Moderate-risk age groups (5-18 years): {risk["age_risk"]["moderate_risk_age_groups"]} cases')
            
            # Add geographic analysis
            doc.add_heading('Geographic Analysis', level=1)
            geo = insights['geographic_analysis']
            
            if 'district_distribution' in geo:
                doc.add_heading('District-wise Case Distribution', level=2)
                for district, count in geo['district_distribution'].items():
                    doc.add_paragraph(f'{district}: {count} cases')
            
            # Add temporal analysis
            doc.add_heading('Temporal Analysis', level=1)
            temporal = insights['temporal_analysis']
            
            if 'weekly_trends' in temporal:
                doc.add_heading('Weekly Case Trends', level=2)
                for week, count in temporal['weekly_trends'].items():
                    doc.add_paragraph(f'Week {week}: {count} cases')
            
            # Add recommendations
            doc.add_heading('Recommendations', level=1)
            self._add_recommendations(doc, insights)
            
            # Save the document
            doc.save(output_path)
            logger.info(f"Report generated successfully at {output_path}")
            
        except Exception as e:
            logger.error(f"Error generating report: {str(e)}")
            raise

    def _add_recommendations(self, doc, insights):
        """Add AI-generated recommendations based on the analysis."""
        try:
            # Analyze high-risk areas
            if 'geographic_analysis' in insights:
                high_risk_districts = insights['geographic_analysis'].get('district_distribution', {})
                if high_risk_districts:
                    doc.add_paragraph("High-risk areas identified:")
                    for district, count in high_risk_districts.items():
                        doc.add_paragraph(f"- {district}: {count} cases")
                    doc.add_paragraph("Recommendations:")
                    doc.add_paragraph("- Increase surveillance in high-risk districts")
                    doc.add_paragraph("- Deploy additional healthcare resources")
                    doc.add_paragraph("- Conduct targeted vaccination campaigns")
            
            # Analyze temporal patterns
            if 'temporal_analysis' in insights:
                weekly_trends = insights['temporal_analysis'].get('weekly_trends', {})
                if weekly_trends:
                    doc.add_paragraph("\nTemporal Pattern Analysis:")
                    doc.add_paragraph("- Monitor weekly trends for early warning signs")
                    doc.add_paragraph("- Prepare for seasonal variations in case numbers")
            
            # Add general recommendations
            doc.add_paragraph("\nGeneral Recommendations:")
            doc.add_paragraph("- Strengthen surveillance systems")
            doc.add_paragraph("- Enhance laboratory capacity")
            doc.add_paragraph("- Improve data collection and reporting")
            doc.add_paragraph("- Conduct regular training for healthcare workers")
            
        except Exception as e:
            logger.error(f"Error adding recommendations: {str(e)}")
            doc.add_paragraph("Error generating recommendations")

def main():
    """Main function to demonstrate usage."""
    try:
        # Initialize the analyzer
        analyzer = ExcelAIAnalyzer()
        
        # Example usage with surveillance data
        excel_file = "example_data/surveillance_data.xlsx"
        output_report = "surveillance_analysis_report.docx"
        
        # Read and analyze data
        df = analyzer.read_excel_data(excel_file)
        insights = analyzer.analyze_data(df)
        
        # Generate report
        analyzer.generate_report(insights, output_report)
        
        logger.info("Analysis completed successfully")
        
    except Exception as e:
        logger.error(f"Error in main execution: {str(e)}")
        raise

if __name__ == "__main__":
    main()